"""
태스크 5 단위 테스트: 행정 서류 자동 생성
- OpenAI API를 mock으로 대체
"""
import pytest
import sys
import os
from io import BytesIO
from unittest.mock import patch, MagicMock
from hypothesis import given, settings
from hypothesis import strategies as st
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from modules.admin_writer import generate_counseling_log, generate_reason_letter, generate_counseling_docx


def make_mock_response(content: str):
    mock_msg = MagicMock()
    mock_msg.content = content
    mock_choice = MagicMock()
    mock_choice.message = mock_msg
    mock_response = MagicMock()
    mock_response.choices = [mock_choice]
    return mock_response


@patch("modules.admin_writer.OpenAI")
def test_counseling_log_generated(mock_openai):
    """상담일지 초안이 생성됨"""
    mock_openai.return_value.chat.completions.create.return_value = make_mock_response(
        "상담일지: [이름] 훈련생은 취업 고민으로 상담을 요청하였음."
    )
    result = generate_counseling_log(
        name="김훈련",
        keywords="취업 고민, 파이썬 기초 부족",
        risk_level="경고",
        absent_info="지각 2회, 조퇴 1회",
    )
    assert len(result) > 0
    assert "[이름]" in result  # PII 마스킹 확인


@patch("modules.admin_writer.OpenAI")
def test_reason_letter_generated(mock_openai):
    """사유서 초안이 생성됨"""
    mock_openai.return_value.chat.completions.create.return_value = make_mock_response(
        "사유서: 2026-04-07 시스템 장애로 인해 출결 체크가 불가하였음."
    )
    result = generate_reason_letter("2026-04-07 시스템 장애로 출결 체크 불가")
    assert len(result) > 0


@patch("modules.admin_writer.OpenAI")
def test_counseling_log_uses_correct_model(mock_openai):
    """gpt-4o-mini 모델 사용 확인"""
    mock_openai.return_value.chat.completions.create.return_value = make_mock_response("초안")
    generate_counseling_log("김훈련", "고민", "정상", "출석 정상")
    call_args = mock_openai.return_value.chat.completions.create.call_args
    assert call_args.kwargs["model"] == "gpt-4o-mini"


def test_no_api_key_counseling_returns_warning(monkeypatch):
    """Property 1: API 키 미설정 시 graceful degradation - counseling log"""
    # Feature: campusguard-deployment, Property 1: API 키 미설정 시 graceful degradation
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    result = generate_counseling_log("김훈련", "고민", "정상", "출석 정상")
    assert isinstance(result, str)
    assert "API" in result or "키" in result


def test_no_api_key_reason_letter_returns_warning(monkeypatch):
    """Property 1: API 키 미설정 시 graceful degradation - reason letter"""
    # Feature: campusguard-deployment, Property 1: API 키 미설정 시 graceful degradation
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    result = generate_reason_letter("시스템 장애")
    assert isinstance(result, str)
    assert "API" in result or "키" in result


@patch("modules.admin_writer.OpenAI")
def test_openai_exception_returns_korean_error(mock_openai, monkeypatch):
    """Property 2: OpenAI 예외 시 한국어 에러 반환"""
    # Feature: campusguard-deployment, Property 2: OpenAI 예외 시 한국어 에러 반환
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    mock_openai.return_value.chat.completions.create.side_effect = Exception("연결 실패")
    result = generate_counseling_log("김훈련", "고민", "정상", "출석 정상")
    assert isinstance(result, str)
    assert "오류" in result or "❌" in result


@patch("modules.admin_writer.OpenAI")
def test_counseling_log_uses_placeholder(mock_openai, monkeypatch):
    """Property 6: 상담일지 이름 플레이스홀더 치환"""
    # Feature: campusguard-deployment, Property 6: 상담일지 이름 플레이스홀더 치환
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    mock_openai.return_value.chat.completions.create.return_value = make_mock_response(
        "상담일지: [이름] 훈련생은 취업 고민으로 상담을 요청하였음."
    )
    result = generate_counseling_log("김훈련", "취업 고민", "경고", "지각 2회")
    assert "[이름]" in result
    assert "김훈련" not in result


# Feature: campusguard-enhancement, Property 11: .docx 변환 필수 필드 포함
@settings(max_examples=100)
@given(
    trainee_name=st.text(min_size=1, alphabet=st.characters(whitelist_categories=("L", "N", "P", "Zs"))),
    course_name=st.text(min_size=1, alphabet=st.characters(whitelist_categories=("L", "N", "P", "Zs"))),
    counseling_date=st.text(min_size=1, alphabet=st.characters(whitelist_categories=("L", "N", "P", "Zs"))),
    content=st.text(min_size=1, alphabet=st.characters(whitelist_categories=("L", "N", "P", "Zs"))),
    action=st.text(min_size=1, alphabet=st.characters(whitelist_categories=("L", "N", "P", "Zs"))),
)
def test_property_11_docx_contains_required_fields(
    trainee_name, course_name, counseling_date, content, action
):
    """
    Property 11: .docx 변환 필수 필드 포함
    Validates: Requirements 4.3, 4.5
    """
    result = generate_counseling_docx(
        trainee_name=trainee_name,
        course_name=course_name,
        counseling_date=counseling_date,
        content=content,
        action=action,
    )

    # 반환값이 비어있지 않아야 함
    assert isinstance(result, bytes)
    assert len(result) > 0

    # python-docx로 파싱하여 필드 값 포함 여부 확인
    doc = Document(BytesIO(result))
    full_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )

    assert trainee_name in full_text
    assert course_name in full_text
    assert counseling_date in full_text
    assert content in full_text
    assert action in full_text


# ── to_docx 단위 테스트 ────────────────────────────────────────────────────────

from modules.admin_writer import to_docx


def test_to_docx_returns_bytes():
    """to_docx는 비어있지 않은 bytes를 반환해야 한다"""
    result = to_docx("테스트 내용입니다.")
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_to_docx_default_title():
    """기본 제목 '상담일지'가 docx에 포함되어야 한다"""
    result = to_docx("내용")
    doc = Document(BytesIO(result))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("상담일지" in h for h in headings)


def test_to_docx_custom_title():
    """커스텀 제목이 docx에 포함되어야 한다"""
    result = to_docx("내용", title="사유서")
    doc = Document(BytesIO(result))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("사유서" in h for h in headings)


def test_to_docx_content_included():
    """본문 텍스트가 docx에 포함되어야 한다"""
    content = "훈련생 김철수는 2024-01-15 상담을 진행하였음."
    result = to_docx(content)
    doc = Document(BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert content in full_text


def test_generate_counseling_docx_returns_bytes():
    """generate_counseling_docx는 비어있지 않은 bytes를 반환해야 한다"""
    result = generate_counseling_docx(
        trainee_name="김훈련",
        course_name="파이썬 기초",
        counseling_date="2024-01-15",
        content="취업 고민 상담",
        action="다음 주 재상담 예정",
    )
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_generate_counseling_docx_contains_fields():
    """generate_counseling_docx 결과 docx에 모든 필드 값이 포함되어야 한다"""
    trainee_name = "이수강"
    course_name = "머신러닝 심화"
    counseling_date = "2024-02-20"
    content = "학습 진도 부진 상담"
    action = "보충 학습 자료 제공"

    result = generate_counseling_docx(
        trainee_name=trainee_name,
        course_name=course_name,
        counseling_date=counseling_date,
        content=content,
        action=action,
    )
    doc = Document(BytesIO(result))
    full_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert trainee_name in full_text
    assert course_name in full_text
    assert counseling_date in full_text
    assert content in full_text
    assert action in full_text
