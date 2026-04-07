"""
태스크 5: 행정 서류 자동 생성
- 상담일지 자동 초안 생성 (HRD-Net 규정 기반)
- 사유서 자동 생성
- .docx 변환 및 HRD-Net 양식 기반 상담일지 생성
"""
import os
import re
from io import BytesIO
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt


def _sanitize_xml(text: str) -> str:
    """XML 비호환 제어 문자를 제거한다 (NULL 바이트 및 C0/C1 제어 문자)."""
    # XML 1.0 허용 문자: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    return re.sub(r"[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]", "", text)

load_dotenv()

COUNSELING_PROMPT = """
당신은 고용노동부 K-디지털 트레이닝 훈련 기관의 행정 담당자입니다.
아래 정보를 바탕으로 공식 상담일지 초안을 작성하세요.

규칙:
- 격식체(~하였음, ~함) 사용
- 훈련생 상태, 상담 내용, 향후 조치 계획 포함
- 200자 내외로 작성
- 개인정보는 [이름], [생년월일] 형태로 표기
"""

REASON_PROMPT = """
당신은 고용노동부 HRD-Net 제출용 사유서를 작성하는 행정 담당자입니다.
아래 상황을 바탕으로 공식 사유서를 작성하세요.

규칙:
- 격식체 사용
- 발생 일시, 사유, 조치 내용 포함
- 150자 내외
"""


def generate_counseling_log(
    name: str,
    keywords: str,
    risk_level: str,
    absent_info: str,
) -> str:
    """
    상담일지 초안 생성.

    Args:
        name: 훈련생 이름 (내부 처리용, 출력 시 [이름]으로 대체)
        keywords: 상담 키워드 (예: "취업 고민, 파이썬 기초 부족")
        risk_level: 위험 등급
        absent_info: 출결 요약 문자열

    Returns:
        상담일지 초안 문자열
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return "⚠️ OpenAI API 키가 설정되지 않아 서류 생성 기능을 사용할 수 없습니다."
    user_input = f"""
훈련생: [이름]
출결 현황: {absent_info}
위험 등급: {risk_level}
상담 키워드: {keywords}
"""
    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": COUNSELING_PROMPT},
                {"role": "user", "content": user_input},
            ],
            max_tokens=400,
            temperature=0.3,
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ 서류 생성 중 오류가 발생했습니다: {str(e)}"


def generate_reason_letter(situation: str) -> str:
    """
    사유서 초안 생성.

    Args:
        situation: 상황 설명 (예: "2026-04-07 시스템 장애로 출결 체크 불가")

    Returns:
        사유서 초안 문자열
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return "⚠️ OpenAI API 키가 설정되지 않아 서류 생성 기능을 사용할 수 없습니다."
    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": REASON_PROMPT},
                {"role": "user", "content": situation},
            ],
            max_tokens=300,
            temperature=0.3,
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ 서류 생성 중 오류가 발생했습니다: {str(e)}"


def to_docx(text: str, title: str = "상담일지") -> bytes:
    """
    텍스트를 .docx 바이트로 변환. python-docx 사용.

    Args:
        text: 문서 본문 텍스트
        title: 문서 제목 (기본값: "상담일지")

    Returns:
        .docx 파일의 bytes
    """
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_counseling_docx(
    trainee_name: str,
    course_name: str,
    counseling_date: str,
    content: str,
    action: str,
) -> bytes:
    """
    HRD-Net 양식 기반 상담일지 .docx 생성.

    Args:
        trainee_name: 훈련생명
        course_name: 과정명
        counseling_date: 상담일시
        content: 상담내용
        action: 조치사항

    Returns:
        .docx 파일의 bytes
    """
    doc = Document()
    doc.add_heading("상담일지", level=1)

    fields = [
        ("훈련생명", _sanitize_xml(trainee_name)),
        ("과정명", _sanitize_xml(course_name)),
        ("상담일시", _sanitize_xml(counseling_date)),
        ("상담내용", _sanitize_xml(content)),
        ("조치사항", _sanitize_xml(action)),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        run = row.cells[0].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
